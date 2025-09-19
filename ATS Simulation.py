import re
from docx import Document

def check_resume(file_path):
    # Keywords to check
    keywords = [
        "Artificial Intelligence",
        "Machine learning",
        "Python flask framework",
        "Hackathon",
        "Github",
        "TensorFlow",
        "PyTorch",
        "Internship",
        "Research",
        "Capstone Project",
        "Open Source",
        "Freelance",
        "Leadership",
        "Teamwork",
        "Problem Solving",
        "Communication",
        "Presentation",
        "Innovation"
    ]
    
    # Step 1: Read the Word file into one big string
    doc = Document(file_path)
    # print(type(doc.paragraphs) )
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    print(text)
    # Step 2: Check each keyword
    missing = []
    for word in keywords:
        if re.search(word, text, re.IGNORECASE):
            print(f"FOUND: {word}")
        else:
            print(f"MISSING: {word}")
            missing.append(word)

    print(missing)
    
    # Step 3: Result
    if not missing:
        print("Resume Passed ✅")
    else:
        print("Resume Rejected ❌")
        print("Missing:", missing)
              

# Example usage
if __name__ == "__main__":
    resume_path = input("Enter path of your resume (.docx): ")
    check_resume(resume_path)



